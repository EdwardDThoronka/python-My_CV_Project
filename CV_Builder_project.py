# To implement the porject, we install pip3 install python-docx using the terminal 
# Import the Document 
from docx import Document
from docx.shared import Inches
document = Document()
# Writing text in the document 
# name = "Edward"
# phone_Number = '00000'
# email = '@gmail.com'
# Here is a profile picture
document.add_picture('me.jpg', width=Inches(2.0))

# Personal information
name = input('What is your name')
phone_Number = input('what is your phone number')
email = input('Enter your email')

document.add_paragraph(
    name + ' | ' + phone_Number + ' | ' + email 
)
# About Me 
document.add_heading('About Me ')
document.add_paragraph(
    input ('Tell about yourself ?')
) 

# Work exprience 
document.add_heading('Work Exprience')
p= document.add_paragraph()

company = input('Enter the company')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + '').bold = True
p.add_run(from_date + '_' + to_date + '\n' ).italic =True

exprience_details =input(
    'Describe your experience at ' + company) 
p.add_run(exprience_details)

# Add more exprience
while True:
    had_more_experience = input(
        'Do you have more experience? Yes or No')
    
    if had_more_experience.lower() == 'yes':
        document.add_heading('Work Exprience')
        p= document.add_paragraph()

        company = input('Enter the company')
        from_date = input('From Date')
        to_date = input('To Date')

        p.add_run(company + '').bold = True
        p.add_run(from_date + '_' + to_date + '\n' ).italic =True

        exprience_details =input(
            'Describe your experience at ' + company + '') 
        p.add_run(exprience_details)
    else:
        break
    
    # Skill 
    document.add_heading('Skills')
    skill = input('Enter skill')
    p= document.add_paragraph(skill)
    p.style='List Bullet'
    
    while True:
        has_more_skills= input('Do you have more skills? Yes or No')
        if has_more_skills.lower() == 'yes':
            skill= input('Enter the Skill')
            p= document.add_paragraph(skill)
            p.style = 'List Bullet'
        else:
            break
    # footer 
    section = document.sections[0]
    footer = section.footer
    p=footer.paragraphs[0]
    p.text= "CV generated using Amigoscode and Intuit QuickBooks course project"
    

# Here is a save function 
document.save('CV.docx')



